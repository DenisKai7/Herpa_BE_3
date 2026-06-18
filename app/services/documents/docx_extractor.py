from io import BytesIO

from docx import Document

from app.core.config import Settings


def extract_docx(data: bytes, settings: Settings) -> dict:
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table_index, table in enumerate(doc.tables, 1):
        parts.append(f"[Tabel {table_index}]")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = "\n".join(parts)
    return {
        "text": text[: settings.max_document_characters],
        "truncated": len(text) > settings.max_document_characters,
        "needs_vision": False,
    }
